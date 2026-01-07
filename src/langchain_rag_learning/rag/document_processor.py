"""Document processing module for parsing and extracting text from various file formats."""

import logging
import mimetypes
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from langchain_rag_learning.core.exceptions import DocumentProcessingError
from langchain_rag_learning.core.models import Document, DocumentChunk

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Main document processor that handles multiple file formats."""
    
    def __init__(self):
        """Initialize the document processor with format-specific parsers."""
        self.parsers = {
            'pdf': PDFParser(),
            'docx': WordParser(),
            'doc': WordParser(),
            'txt': TextParser(),
            'md': MarkdownParser(),
            'markdown': MarkdownParser(),
            'html': HTMLParser(),
            'htm': HTMLParser()
        }
        self.metadata_extractor = MetadataExtractor()
        self.text_cleaner = TextCleaner()
    
    async def process_document(self, file_path: str, document: Document) -> Tuple[str, Dict[str, Any]]:
        """
        Process a document and extract text content and metadata.
        
        Args:
            file_path: Path to the document file
            document: Document model instance
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            DocumentProcessingError: If processing fails
        """
        try:
            # Determine file type
            file_type = self._get_file_type(file_path, document.file_type)
            
            # Get appropriate parser
            parser = self.parsers.get(file_type)
            if not parser:
                raise DocumentProcessingError(f"No parser available for file type: {file_type}")
            
            # Extract text content
            logger.info(f"Processing document: {document.filename} (type: {file_type})")
            raw_text = await parser.extract_text(file_path)
            
            # Clean and preprocess text
            cleaned_text = self.text_cleaner.clean_text(raw_text)
            
            # Extract metadata
            metadata = await self.metadata_extractor.extract_metadata(
                file_path, document, cleaned_text
            )
            
            logger.info(f"Successfully processed document: {document.filename}")
            return cleaned_text, metadata
            
        except Exception as e:
            error_msg = f"Failed to process document {document.filename}: {str(e)}"
            logger.error(error_msg)
            raise DocumentProcessingError(error_msg) from e
    
    def _get_file_type(self, file_path: str, declared_type: str) -> str:
        """
        Determine the actual file type from path and declared type.
        
        Args:
            file_path: Path to the file
            declared_type: Declared file type from document model
            
        Returns:
            Normalized file type string
        """
        # First try the declared type
        if declared_type and declared_type.lower() in self.parsers:
            return declared_type.lower()
        
        # Fall back to file extension
        path = Path(file_path)
        extension = path.suffix.lower().lstrip('.')
        
        if extension in self.parsers:
            return extension
        
        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type:
            type_mapping = {
                'application/pdf': 'pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                'application/msword': 'doc',
                'text/plain': 'txt',
                'text/markdown': 'md',
                'text/html': 'html'
            }
            return type_mapping.get(mime_type, 'txt')
        
        # Default to text
        return 'txt'


class BaseParser:
    """Base class for document parsers."""
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text content from a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If extraction fails
        """
        raise NotImplementedError("Subclasses must implement extract_text method")


class PDFParser(BaseParser):
    """Parser for PDF documents."""
    
    async def extract_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
            
            return "\n\n".join(text_content)
            
        except ImportError:
            raise DocumentProcessingError("PyPDF2 library not installed. Install with: pip install PyPDF2")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse PDF: {str(e)}")


class WordParser(BaseParser):
    """Parser for Word documents (.docx, .doc)."""
    
    async def extract_text(self, file_path: str) -> str:
        """Extract text from Word document."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    text_content.append("\n--- Table ---\n" + "\n".join(table_text))
            
            return "\n\n".join(text_content)
            
        except ImportError:
            raise DocumentProcessingError("python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse Word document: {str(e)}")


class TextParser(BaseParser):
    """Parser for plain text files."""
    
    async def extract_text(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError(f"Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse text file: {str(e)}")


class MarkdownParser(BaseParser):
    """Parser for Markdown files."""
    
    async def extract_text(self, file_path: str) -> str:
        """Extract text from Markdown file."""
        try:
            # First extract as plain text
            text_parser = TextParser()
            raw_content = await text_parser.extract_text(file_path)
            
            # Basic markdown processing - remove common markdown syntax
            # but preserve structure and content
            processed_content = self._process_markdown(raw_content)
            
            return processed_content
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse Markdown file: {str(e)}")
    
    def _process_markdown(self, content: str) -> str:
        """Process markdown content to clean up syntax while preserving structure."""
        lines = content.split('\n')
        processed_lines = []
        
        for line in lines:
            # Convert headers to plain text with structure indicators
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('# ').strip()
                if header_text:
                    processed_lines.append(f"{'=' * level} {header_text} {'=' * level}")
                continue
            
            # Remove markdown formatting but keep content
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
            line = re.sub(r'\*(.*?)\*', r'\1', line)      # Italic
            line = re.sub(r'`(.*?)`', r'\1', line)        # Inline code
            line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line)  # Links
            
            # Handle lists
            if re.match(r'^\s*[-*+]\s+', line):
                line = re.sub(r'^\s*[-*+]\s+', '• ', line)
            elif re.match(r'^\s*\d+\.\s+', line):
                line = re.sub(r'^\s*\d+\.\s+', '• ', line)
            
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)


class HTMLParser(BaseParser):
    """Parser for HTML files."""
    
    async def extract_text(self, file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text with some structure preservation
            text_content = []
            
            # Extract title
            title = soup.find('title')
            if title and title.get_text().strip():
                text_content.append(f"=== {title.get_text().strip()} ===")
            
            # Extract headings with structure
            for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                level = int(heading.name[1])
                heading_text = heading.get_text().strip()
                if heading_text:
                    text_content.append(f"{'=' * level} {heading_text} {'=' * level}")
            
            # Extract paragraphs and other text content
            for element in soup.find_all(['p', 'div', 'article', 'section']):
                text = element.get_text().strip()
                if text and len(text) > 10:  # Filter out very short text
                    text_content.append(text)
            
            return '\n\n'.join(text_content)
            
        except ImportError:
            raise DocumentProcessingError("beautifulsoup4 library not installed. Install with: pip install beautifulsoup4")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to parse HTML file: {str(e)}")


class MetadataExtractor:
    """Extracts metadata from documents and files."""
    
    async def extract_metadata(self, file_path: str, document: Document, text_content: str) -> Dict[str, Any]:
        """
        Extract comprehensive metadata from document.
        
        Args:
            file_path: Path to the document file
            document: Document model instance
            text_content: Extracted text content
            
        Returns:
            Dictionary containing extracted metadata
        """
        metadata = {}
        
        try:
            path = Path(file_path)
            
            # Basic file metadata
            metadata.update({
                'file_size_bytes': document.file_size,
                'file_extension': path.suffix.lower(),
                'mime_type': document.mime_type,
                'original_filename': document.original_filename
            })
            
            # Text analysis metadata
            metadata.update(self._analyze_text_content(text_content))
            
            # File-specific metadata
            file_metadata = await self._extract_file_specific_metadata(file_path, document.file_type)
            metadata.update(file_metadata)
            
        except Exception as e:
            logger.warning(f"Failed to extract some metadata: {e}")
        
        return metadata
    
    def _analyze_text_content(self, text: str) -> Dict[str, Any]:
        """Analyze text content and extract statistics."""
        if not text:
            return {}
        
        lines = text.split('\n')
        words = text.split()
        
        return {
            'character_count': len(text),
            'word_count': len(words),
            'line_count': len(lines),
            'paragraph_count': len([line for line in lines if line.strip()]),
            'average_words_per_line': len(words) / max(len(lines), 1),
            'language_detected': self._detect_language(text[:1000])  # Sample first 1000 chars
        }
    
    def _detect_language(self, text_sample: str) -> str:
        """Simple language detection based on character patterns."""
        # Basic language detection - can be enhanced with proper libraries
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text_sample))
        total_chars = len(text_sample.replace(' ', ''))
        
        if total_chars > 0 and chinese_chars / total_chars > 0.3:
            return 'zh'
        return 'en'  # Default to English
    
    async def _extract_file_specific_metadata(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Extract file-type specific metadata."""
        metadata = {}
        
        try:
            if file_type == 'pdf':
                metadata.update(await self._extract_pdf_metadata(file_path))
            elif file_type in ['docx', 'doc']:
                metadata.update(await self._extract_word_metadata(file_path))
            elif file_type == 'html':
                metadata.update(await self._extract_html_metadata(file_path))
        except Exception as e:
            logger.warning(f"Failed to extract {file_type} specific metadata: {e}")
        
        return metadata
    
    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF-specific metadata."""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'page_count': len(pdf_reader.pages),
                    'pdf_version': getattr(pdf_reader, 'pdf_header', 'Unknown')
                }
                
                # Extract PDF metadata if available
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.update({
                        'pdf_title': pdf_meta.get('/Title', ''),
                        'pdf_author': pdf_meta.get('/Author', ''),
                        'pdf_subject': pdf_meta.get('/Subject', ''),
                        'pdf_creator': pdf_meta.get('/Creator', ''),
                        'pdf_producer': pdf_meta.get('/Producer', ''),
                        'pdf_creation_date': str(pdf_meta.get('/CreationDate', '')),
                        'pdf_modification_date': str(pdf_meta.get('/ModDate', ''))
                    })
                
                return metadata
                
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata: {e}")
            return {}
    
    async def _extract_word_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract Word document metadata."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'doc_title': props.title or '',
                    'doc_author': props.author or '',
                    'doc_subject': props.subject or '',
                    'doc_keywords': props.keywords or '',
                    'doc_comments': props.comments or '',
                    'doc_created': str(props.created) if props.created else '',
                    'doc_modified': str(props.modified) if props.modified else ''
                })
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Failed to extract Word metadata: {e}")
            return {}
    
    async def _extract_html_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract HTML metadata."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            metadata = {}
            
            # Extract meta tags
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                name = tag.get('name') or tag.get('property')
                content = tag.get('content')
                if name and content:
                    metadata[f'meta_{name}'] = content
            
            # Extract title
            title = soup.find('title')
            if title:
                metadata['html_title'] = title.get_text().strip()
            
            # Count elements
            metadata.update({
                'heading_count': len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])),
                'paragraph_count': len(soup.find_all('p')),
                'link_count': len(soup.find_all('a')),
                'image_count': len(soup.find_all('img'))
            })
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Failed to extract HTML metadata: {e}")
            return {}


class TextCleaner:
    """Cleans and preprocesses extracted text content."""
    
    def clean_text(self, text: str) -> str:
        """
        Clean and preprocess text content.
        
        Args:
            text: Raw text content
            
        Returns:
            Cleaned text content
        """
        if not text:
            return ""
        
        # Normalize whitespace
        text = self._normalize_whitespace(text)
        
        # Remove unwanted characters
        text = self._remove_unwanted_characters(text)
        
        # Fix common encoding issues
        text = self._fix_encoding_issues(text)
        
        # Normalize line breaks
        text = self._normalize_line_breaks(text)
        
        return text.strip()
    
    def _normalize_whitespace(self, text: str) -> str:
        """Normalize whitespace characters."""
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace tabs with spaces
        text = text.replace('\t', ' ')
        
        # Remove trailing whitespace from lines
        lines = text.split('\n')
        lines = [line.rstrip() for line in lines]
        
        return '\n'.join(lines)
    
    def _remove_unwanted_characters(self, text: str) -> str:
        """Remove unwanted characters and control characters."""
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Remove zero-width characters
        text = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)
        
        return text
    
    def _fix_encoding_issues(self, text: str) -> str:
        """Fix common encoding issues."""
        # Common encoding fixes
        replacements = {
            'â€™': "'",
            'â€œ': '"',
            'â€': '"',
            'â€"': '—',
            'â€"': '–',
            'Â': ' ',
            'â€¦': '...'
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text
    
    def _normalize_line_breaks(self, text: str) -> str:
        """Normalize line breaks and paragraph spacing."""
        # Replace multiple consecutive newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Ensure proper paragraph separation
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text